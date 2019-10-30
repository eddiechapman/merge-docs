"""
merge_docs.py

Concatenate the mission and skills documents for each degree.

Create a 'description' document is created for degrees that have a valid
skills and valid mission document. The resulting document combines the two.

Incomplete, invalid, course, and unknown documents are copied to new
directories for inspection.

The contents of the input directory are moved.

                        [ID]-[Category].docx

    - [ID] a three digit number (zero-padded) representing a degree
    - [Category] one of "skills", "courses", or "mission"

Output directories:

    - descriptions: new description documents
    - incomplete: mission and skills documents that could not be paired
    - courses: courses documents
    - error: files that had no text or otherwise caused an error
    - unknown: files from the input directory that could not be sorted

This script is part of the PERVADE project for the data science
curriculum sub-project.

"""
import argparse
import logging
import pathlib
import re
import docx


def main(args):
    logging.basicConfig(
        level=args.log_level or logging.INFO,
        filename='app.log',
        format='%(levelname)s:%(asctime)s:%(message)s'
    )

    INPUT = pathlib.Path(args.input)
    OUTPUT = pathlib.Path(args.output)
    OUTPUT.mkdir(exist_ok=True)

    DESCRIPTIONS = OUTPUT / 'descriptions'
    INCOMPLETE = OUTPUT / 'incomplete'
    COURSES = OUTPUT / 'courses'
    ERROR = OUTPUT / 'error'
    UNKNOWN = OUTPUT / 'unknown'

    DESCRIPTIONS.mkdir(exist_ok=True)
    INCOMPLETE.mkdir(exist_ok=True)
    COURSES.mkdir(exist_ok=True)
    ERROR.mkdir(exist_ok=True)
    UNKNOWN.mkdir(exist_ok=True)

    # Sort the document paths by degree ID and category
    documents = dict()
    pattern = r'(?P<id>\d{3})-(?P<category>skills|courses|mission).docx'
    for f in INPUT.glob('*.docx'):
        logging.debug(f'Sorting file: {f}')
        m = re.search(pattern, f.name, re.IGNORECASE)
        if m:
            logging.debug(f'{f} matched: {m.groupdict()}')
            if m.group('id') not in documents:
                documents[m.group('id')] = dict()
            documents[m.group('id')][m.group('category').lower()] = f

    # Process sorted documents
    for degree, docs in documents.items():
        skills = docs.get('skills')
        mission = docs.get('mission')
        courses = docs.get('courses')

        # Course documents are saved for archiving
        if courses:
            logging.debug(f'Moving {courses} to {COURSES}')
            courses.rename(COURSES / courses.name)

        # Paired mission and skills documents are combined
        if mission and skills:
            m_doc = docx.Document(mission)
            s_doc = docx.Document(skills)
            m_txt = '\n'.join([p.text for p in m_doc.paragraphs])
            s_txt = '\n'.join([p.text for p in s_doc.paragraphs])

            # Both documents are present and contain text
            if m_txt and s_txt:
                path = DESCRIPTIONS / f'{degree}-description.docx'
                description = docx.Document()
                description.add_paragraph(m_txt)
                description.add_paragraph(s_txt)
                description.save(path)
                logging.debug(f'Description for #{degree} written to {path}')

            # Both documents are present but one or both is blank
            else:
                if not m_txt:
                    logging.error(f'Bad document: {mission.name}')
                    mission.rename(ERROR / mission.name)
                if not s_txt:
                    logging.error(f'Bad document {skills.name}')
                    skills.rename(ERROR / skills.name)

        # Incomplete document sets
        else:
            if mission:
                mission.rename(INCOMPLETE / mission.name)
                logging.debug(f'Moving incomplete: {mission.name}')
            elif skills:
                skills.rename(INCOMPLETE / skills.name)
                logging.debug(f'Moving incomplete: {skills.name}')

    # Separate unknown files from input directory
    known = set()
    for docs in documents.values():
        known.update(docs.values())

    for doc in INPUT.iterdir():
        if not doc in known:
            doc.rename(UNKNOWN / doc.name)
            logging.warn(f'Moving unknown file: {doc.name}')


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description=__doc__)
    group = parser.add_mutually_exclusive_group()
    group.add_argument(
        '-v', '--verbose',
        help='Verbose (debug) logging level.',
        const=logging.DEBUG,
        dest='log_level',
        nargs='?',
    )
    group.add_argument(
        '-q', '--quiet',
        help='Silent mode, only log warnings and errors.',
        const=logging.WARN,
        dest='log_level',
        nargs='?',
    )
    parser.add_argument(
        '-i', '--input',
        help=(f'Path to directory of .docx files in format: '
              f'###-[skills|mission|courses].docx'),
        metavar='input',
        type=str
    )
    parser.add_argument(
        '-o', '--output',
        help=(f'Path to the directory where the resulting files '
              f'will be written.'),
        metavar='output',
        default='./output',
        type=str
    )
    args = parser.parse_args()
    main(args)



